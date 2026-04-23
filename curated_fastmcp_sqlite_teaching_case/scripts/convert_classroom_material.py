#!/usr/bin/env python3
from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "classroom_material.md"
OUTPUT_DOCX = ROOT / "docs" / "classroom_material.docx"
OUTPUT_PDF = ROOT / "docs" / "classroom_material.pdf"


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    current = lines[index].strip()
    next_line = lines[index + 1].strip()
    return current.startswith("|") and current.endswith("|") and re.match(r"^\|[\s:\-|]+\|$", next_line) is not None


def collect_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    table_lines = []
    while index < len(lines):
        line = lines[index].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        table_lines.append(line)
        index += 1

    rows = []
    for line_number, line in enumerate(table_lines):
        if line_number == 1 and re.match(r"^\|[\s:\-|]+\|$", line):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
    return rows, index


def split_inline_code(text: str) -> list[tuple[str, bool]]:
    parts = re.split(r"(`[^`]+`)", text)
    result = []
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            result.append((part[1:-1], True))
        else:
            result.append((part, False))
    return result


def add_docx_inline(paragraph, text: str) -> None:
    for chunk, is_code in split_inline_code(text):
        run = paragraph.add_run(chunk)
        if is_code:
            run.font.name = "Menlo"
            run.font.size = Pt(9)


def add_docx_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for column_index in range(column_count):
            text = row[column_index] if column_index < len(row) else ""
            cells[column_index].text = text
            for paragraph in cells[column_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True


def build_docx(markdown: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.size = Pt(22)

    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                for code_index, code_line in enumerate(code_lines):
                    run = paragraph.add_run(code_line)
                    run.font.name = "Menlo"
                    run.font.size = Pt(8.5)
                    if code_index < len(code_lines) - 1:
                        run.add_break(WD_BREAK.LINE)
                in_code = False
                code_lines = []
            else:
                in_code = True
                code_lines = []
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if is_table_start(lines, index):
            rows, index = collect_table(lines, index)
            add_docx_table(document, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = heading.group(2)
            if level == 1:
                paragraph = document.add_paragraph(text, style="Title")
            else:
                paragraph = document.add_heading(text, level=level)
            paragraph.paragraph_format.space_after = Pt(8)
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_docx_inline(paragraph, bullet.group(1))
        elif numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_docx_inline(paragraph, numbered.group(1))
        else:
            paragraph = document.add_paragraph()
            add_docx_inline(paragraph, stripped)
        index += 1

    document.save(OUTPUT_DOCX)


def escape_pdf_text(text: str) -> str:
    escaped = html.escape(text)
    return re.sub(r"`([^`]+)`", r'<font name="Courier">\1</font>', escaped)


def add_pdf_table(story: list, rows: list[list[str]], styles: dict[str, ParagraphStyle]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = []
    for row in rows:
        normalized.append(row + [""] * (column_count - len(row)))

    page_width = A4[0] - 36 * mm
    col_width = page_width / column_count
    data = [
        [Paragraph(escape_pdf_text(cell), styles["table_header" if row_index == 0 else "table_cell"]) for cell in row]
        for row_index, row in enumerate(normalized)
    ]
    table = Table(data, colWidths=[col_width] * column_count, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111827")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 6))


def build_pdf(markdown: str) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    base = getSampleStyleSheet()
    styles = {
        "h1": ParagraphStyle(
            "h1",
            parent=base["Heading1"],
            fontName="STSong-Light",
            fontSize=18,
            leading=24,
            spaceAfter=10,
            alignment=TA_LEFT,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "h2",
            parent=base["Heading2"],
            fontName="STSong-Light",
            fontSize=14,
            leading=19,
            spaceBefore=8,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "h3": ParagraphStyle(
            "h3",
            parent=base["Heading3"],
            fontName="STSong-Light",
            fontSize=12,
            leading=17,
            spaceBefore=6,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "body",
            parent=base["BodyText"],
            fontName="STSong-Light",
            fontSize=10,
            leading=15,
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "bullet",
            parent=base["BodyText"],
            fontName="STSong-Light",
            fontSize=10,
            leading=15,
            leftIndent=14,
            firstLineIndent=-9,
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "code": ParagraphStyle(
            "code",
            parent=base["Code"],
            fontName="STSong-Light",
            fontSize=7.5,
            leading=10,
            leftIndent=4,
            rightIndent=4,
            backColor=colors.HexColor("#F6F8FA"),
            borderColor=colors.HexColor("#D0D7DE"),
            borderWidth=0.3,
            borderPadding=5,
            wordWrap="CJK",
        ),
        "table_header": ParagraphStyle(
            "table_header",
            fontName="STSong-Light",
            fontSize=8,
            leading=10,
            wordWrap="CJK",
        ),
        "table_cell": ParagraphStyle(
            "table_cell",
            fontName="STSong-Light",
            fontSize=7.5,
            leading=9.5,
            wordWrap="CJK",
        ),
    }

    story = []
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                story.append(Preformatted("\n".join(code_lines), styles["code"]))
                story.append(Spacer(1, 5))
                in_code = False
                code_lines = []
            else:
                in_code = True
                code_lines = []
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if is_table_start(lines, index):
            rows, index = collect_table(lines, index)
            add_pdf_table(story, rows, styles)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            style_name = "h1" if level == 1 else "h2" if level == 2 else "h3"
            if level == 2 and story:
                story.append(Spacer(1, 3))
            story.append(Paragraph(escape_pdf_text(heading.group(2)), styles[style_name]))
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if bullet:
            story.append(Paragraph(f"• {escape_pdf_text(bullet.group(1))}", styles["bullet"]))
        elif numbered:
            story.append(Paragraph(f"{numbered.group(1)}. {escape_pdf_text(numbered.group(2))}", styles["bullet"]))
        else:
            story.append(Paragraph(escape_pdf_text(stripped), styles["body"]))
        index += 1

    document = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="教材文件：用 FastMCP 與 SQLite 教 AI Agent 判讀空壓機資料",
    )
    document.build(story)


def main() -> None:
    markdown = SOURCE_MD.read_text(encoding="utf-8")
    build_docx(markdown)
    build_pdf(markdown)
    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Wrote {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
